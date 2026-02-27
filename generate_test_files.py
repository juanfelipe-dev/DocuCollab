"""
Script to generate 5 Word documents and 5 Excel spreadsheets for testing DocuCollab
"""

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
import os

# Create media/documents directory if it doesn't exist
docs_dir = "media/documents"
os.makedirs(docs_dir, exist_ok=True)

print("\n" + "="*60)
print("🎯 Creating Test Documents for DocuCollab")
print("="*60 + "\n")

# Create 5 Word Documents
print("📝 Creating Word Documents (.docx):\n")
for i in range(1, 6):
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Test Word Document {i}', 0)
    
    # Add content
    doc.add_paragraph(f'\nThis is a sample Word document for testing DocuCollab.')
    doc.add_paragraph(f'\nDocument Number: {i}')
    doc.add_paragraph(f'\nDescription: This document contains sample content to test the collaborative editing features of DocuCollab.')
    
    # Add a table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Status'
    
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True
    
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Document Upload'
    row_cells[1].text = '✅ Working'
    
    row_cells = table.rows[2].cells
    row_cells[0].text = 'Document Sharing'
    row_cells[1].text = '✅ Working'
    
    # Add footer
    doc.add_paragraph('\n' + '-'*50)
    doc.add_paragraph('Generated for DocuCollab Testing - Sample Document')
    
    filename = f'{docs_dir}/Test_Word_Document_{i}.docx'
    doc.save(filename)
    print(f"  ✅ Created: Test_Word_Document_{i}.docx")

print("\n📊 Creating Excel Spreadsheets (.xlsx):\n")

# Create 5 Excel Spreadsheets
for i in range(1, 6):
    wb = Workbook()
    ws = wb.active
    ws.title = f"Sheet_{i}"
    
    # Add headers
    headers = ['Item', 'Quantity', 'Price', 'Total']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col)
        cell.value = header
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    
    # Add sample data
    data = [
        ['Product A', 10, 25.50, 255.00],
        ['Product B', 15, 30.00, 450.00],
        ['Product C', 8, 45.75, 366.00],
        ['Product D', 20, 15.25, 305.00],
        ['Product E', 12, 55.00, 660.00],
    ]
    
    for row_idx, row_data in enumerate(data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.value = value
            if col_idx >= 3:  # Format price columns
                cell.number_format = '$#,##0.00'
    
    # Add summary row
    summary_row = len(data) + 3
    ws.cell(row=summary_row, column=1).value = "TOTAL"
    ws.cell(row=summary_row, column=1).font = Font(bold=True)
    
    total_cell = ws.cell(row=summary_row, column=4)
    total_cell.value = f"=SUM(D2:D{len(data)+1})"
    total_cell.font = Font(bold=True)
    total_cell.number_format = '$#,##0.00'
    
    # Adjust column widths
    ws.column_dimensions['A'].width = 15
    ws.column_dimensions['B'].width = 12
    ws.column_dimensions['C'].width = 12
    ws.column_dimensions['D'].width = 12
    
    # Add title row at top (shift data down)
    ws.insert_rows(1)
    title_cell = ws.cell(row=1, column=1)
    title_cell.value = f"Test Excel Spreadsheet {i}"
    title_cell.font = Font(bold=True, size=14)
    
    filename = f'{docs_dir}/Test_Excel_Spreadsheet_{i}.xlsx'
    wb.save(filename)
    print(f"  ✅ Created: Test_Excel_Spreadsheet_{i}.xlsx")

print("\n" + "="*60)
print("✅ All Test Documents Created Successfully!")
print("="*60)
print(f"\n📁 Location: {os.path.abspath(docs_dir)}\n")
print("📋 Summary:")
print("  • 5 Word Documents (.docx)")
print("  • 5 Excel Spreadsheets (.xlsx)")
print("  • Total: 10 test files ready\n")
print("🚀 Next Steps:")
print("  1. Login to http://localhost:8000")
print("  2. Go to Documents section")
print("  3. Click '+ New Document'")
print("  4. Upload any of these test files")
print("  5. Test sharing, viewing, and editing features\n")
print("="*60 + "\n")
